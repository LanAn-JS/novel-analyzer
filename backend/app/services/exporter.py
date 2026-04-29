"""Novel Analyzer - 导出服务"""

import json
import os
from datetime import datetime
from typing import Optional
from pathlib import Path

from ..config import APP_SIGNATURE, APP_NAME


def analysis_to_markdown(novel_title: str, analysis: dict) -> str:
    """将分析结果转为Markdown"""
    md = f"""# 《{novel_title}》分析报告

> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
> 生成工具：{APP_NAME}  
> {APP_SIGNATURE}

---

## 📋 标签

**题材标签：** {', '.join(analysis.get('genre_tags', [])) or '无'}

**风格标签：** {', '.join(analysis.get('style_tags', [])) or '无'}

---

## 📝 内容总结

{analysis.get('summary', '未分析')}

---

## ✍️ 文笔分析

{analysis.get('writing_style', '未分析')}

---

## 🏗️ 结构分析

{analysis.get('structure', '未分析')}

---

## 📑 大纲

{analysis.get('outline', '未分析')}

---

## 📑 细纲

{analysis.get('detailed_outline', '未分析')}

---

## 🌍 世界观设定

{analysis.get('worldview', '未分析')}

---

## 👤 人物分析

{analysis.get('character_growth', '未分析')}

---

## 🔥 流行性分析

{analysis.get('popularity_analysis', '未分析')}

**题材热度：** {analysis.get('genre_heat', 0)}/10  
**上升潜力：** {analysis.get('rise_potential', 0)}/10

---

*{APP_SIGNATURE}*
"""
    return md


def analysis_to_json(novel_title: str, novel_info: dict, analysis: dict) -> str:
    """将分析结果转为JSON"""
    data = {
        "app": APP_NAME,
        "signature": APP_SIGNATURE,
        "generated_at": datetime.now().isoformat(),
        "novel": {
            "title": novel_title,
            **novel_info,
        },
        "analysis": analysis,
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


async def export_analysis(
    novel_title: str,
    novel_info: dict,
    analysis: dict,
    fmt: str = "markdown",
    output_path: Optional[str] = None,
) -> str:
    """
    导出分析报告
    返回文件路径
    """
    if output_path is None:
        from ..config import DATA_DIR
        safe_title = "".join(c if c.isalnum() or c in "_ -" else "_" for c in novel_title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = DATA_DIR / "exports"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = str(output_dir / f"{safe_title}_{timestamp}")

    if fmt == "markdown" or fmt == "md":
        content = analysis_to_markdown(novel_title, analysis)
        path = output_path + ".md"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    elif fmt == "json":
        content = analysis_to_json(novel_title, novel_info, analysis)
        path = output_path + ".json"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    elif fmt == "pdf":
        # Markdown → PDF (via weasyprint)
        md_content = analysis_to_markdown(novel_title, analysis)
        import markdown as md_lib
        html_content = md_lib.markdown(md_content, extensions=["tables", "fenced_code"])
        
        # 添加CSS样式
        html_full = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: "Noto Sans SC", "Microsoft YaHei", sans-serif; padding: 40px; line-height: 1.8; }}
h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
h2 {{ color: #34495e; margin-top: 30px; }}
blockquote {{ background: #f8f9fa; border-left: 4px solid #3498db; padding: 10px 20px; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
</style>
</head>
<body>{html_content}</body>
</html>"""
        
        path = output_path + ".pdf"
        from weasyprint import HTML
        HTML(string=html_full).write_pdf(path)
        return path

    elif fmt == "docx":
        # Markdown → DOCX
        md_content = analysis_to_markdown(novel_title, analysis)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        doc.add_heading(f"《{novel_title}》分析报告", level=0)
        
        # 简单按行转换
        for line in md_content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.style = doc.styles['Intense Quote']
            elif line.startswith("---"):
                doc.add_paragraph("─" * 40)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)
        
        path = output_path + ".docx"
        doc.save(path)
        return path

    else:
        raise ValueError(f"不支持的导出格式: {fmt}")
